import pandas as pd
import subprocess
import re
import docx
from docx.shared import Inches

class Samplesheet:
    def __init__(self, folder_out: str):
        self.folder_out = folder_out

    def listCsvFiles(self) -> list[str]:
        command = f"ls $PWD/*.csv"
        return subprocess.run(command, 
                              shell=True, 
                              capture_output=True,
                              text=True, 
                              check=True).stdout.split('\n')

    @property
    def sample_table(self):
        list_file = self.listCsvFiles()
        collector = []
        for file in list_file:
            try:
                rs = re.search(r'Tab_S(\d+)_([^_]+)_.*_([^_]+).csv', file.split('/')[-1])
                collector.append(rs.groups() + (file,))
            except:
                continue
        
        data = pd.DataFrame(collector, columns=['index', 'prefix', 'pros', 'file'])
        data['type'] = data[['prefix', 'pros']].apply(lambda x: self.Categorize(x['prefix'], x['pros']), axis='columns')
        data['index'] = data['index'].astype(int)
        data = data.sort_values(by=['type', 'index']).reset_index(drop=True)
        return data
    
    @staticmethod
    def Categorize(prefix:str, pros: str):
        if re.search(r'[A-Z]', pros):
            return 'prs'
        if prefix == 'mean':
            return 'mean_r2'
        if prefix == 'coverage':
            return 'coverage'
        else:
            return 'percentile'

           
class PdfPrinter:
    def __init__(self, docx_path: str, path: str):
        self.docx_path = docx_path
        self.docx = docx.Document()
        self.ss = Samplesheet(path)
        self.samplesheet = self.ss.sample_table
        self.caption_libs = {
            'mean_r2' : 'Imputation accuracy (mean and standard deviation across 22 autosomes) for eight genotyping arrays and six LPS coverages, evaluated across five populations for variant with allel frequency ',
            'coverage' : 'Imputation coverage (mean and standard deviation across 22 autosomes) for eight genotyping arrays and six LPS coverages, evaluated across five populations for variant with allel frequency ',
            'prs': 'Mean and the standard deviation of PGS correlation of eight genotyping arrays and six LPS coverages of the phenotype ',
            'percentile': 'Mean absolute difference of percentile ranking between PGSs estimated from imputed genotyping data of eight genotyping arrays and six LPS coverages and PGS estimated from WGS in 5 different populations with PRsice p-value setting of '
        }

        self.phenotype_lib = {
            'METABOLIC': 'metabolic',
            'HEIGHT': 'height',
            'DIABETES': 'diabetes',
            'BMI': 'the phenotype body mass index (BMI)'
        }


    def rendering(self):
        ## Render mean r^2
        mean_r2_table = self.samplesheet[self.samplesheet['type'] == 'mean_r2']
        for row in mean_r2_table.to_dict(orient='records'):
            caption = f'Table S. {row["index"]} ' + self.caption_libs[row['type']] + row['pros']
            self.addCaption(caption=caption)
            self.addTable(row['file'], row['type'])

        ## Render coverage
        mean_r2_table = self.samplesheet[self.samplesheet['type'] == 'coverage']
        for row in mean_r2_table.to_dict(orient='records'):
            caption = f'Table S. {row["index"]} ' + self.caption_libs[row['type']] + row['pros']
            self.addCaption(caption=caption)
            self.addTable(row['file'], row['type'])

        ## Render prs
        mean_r2_table = self.samplesheet[self.samplesheet['type'] == 'prs']
        for row in mean_r2_table.to_dict(orient='records'):
            caption = f'Table S. {row["index"]} ' + self.caption_libs[row['type']] + self.phenotype_lib[row['pros']]
            self.addCaption(caption=caption)
            self.addTable(row['file'], row['type'])

        ## Render percentile
        mean_r2_table = self.samplesheet[self.samplesheet['type'] == 'percentile']
        for row in mean_r2_table.to_dict(orient='records'):
            caption = f'Table S. {row["index"]} ' + self.caption_libs[row['type']] + row['pros']
            self.addCaption(caption=caption)
            self.addTable(row['file'], row['type'])
        
        ## Save to file
        self.save()

    def addCaption(self, caption: str):
        p = self.docx.add_paragraph()
        p.add_run(caption)
    
    def addTable(self, table_path: str, type: str):
        data = pd.read_csv(table_path)
        t = self.docx.add_table(data.shape[0]+1, data.shape[1])

        # correct column name.
        if type == 'percentile':
            data = data.rename(columns={'array': 'Array/LPS', 'trait': 'Trait'})
      
        # add the header rows.
        for j in range(data.shape[-1]):
            t.cell(0,j).text = data.columns[j]

        # add the rest of the data frame
        for i in range(data.shape[0]):
            for j in range(data.shape[-1]):
                t.cell(i+1,j).text = str(data.values[i,j])
        
        # set style
        t.style = 'Table Grid'

        # adjust width
        for index, cell in enumerate(t.rows[0].cells):
            if index == 0:
                continue
            cell.width = Inches(5)

        # break page
        self.docx.add_page_break()

    def save(self) -> None:
        self.docx.save(self.docx_path)

            
    

if __name__ == '__main__':
    path = "/home/ktest/project/truongphi/RND/lps_paper/out_tables"
    ss = Samplesheet(path)

    # print(ss.listCsvFiles())
    # print(ss.sample_table)

    ## Render words
    # pp = PdfPrinter('test.docx')

    # pp.addTable("/home/ktest/project/truongphi/RND/lps_paper/out_tables/Tab_S3_mean_r2_summary_bin_(0.05–0.5].csv")

    # pp.save()

    pp = PdfPrinter('supplements.docx', path)
    pp.rendering()

    