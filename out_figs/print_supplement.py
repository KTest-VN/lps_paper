import pandas as pd
import subprocess
import re
import docx
from docx.shared import Inches
import fitz
import tempfile
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Samplesheet:
    def __init__(self, folder_out: str):
        self.folder_out = folder_out

    def listPdfFiles(self) -> list[str]:
        command = f"ls $PWD/Fig_S*.pdf"
        return subprocess.run(command, 
                              shell=True, 
                              capture_output=True,
                              text=True, 
                              check=True).stdout.split('\n')

    @property
    def sample_table(self):
        list_file = self.listPdfFiles()
        collector = []
        for file in list_file:
            try:
                rs = re.search(r'Fig_S(\d+)_Pt_(.*).pdf', file.split('/')[-1])
                collector.append(rs.groups() + (file,))
            except:
                continue
        
        data = pd.DataFrame(collector, columns=['index', 'percentile', 'file'])
        data['index'] = data['index'].astype(int)
        data = data.sort_values(by=['index'], ascending=True)
        return data
           
class PdfPrinter:
    def __init__(self, docx_path: str, path: str):
        self.docx_path = docx_path
        self.docx = docx.Document()
        self.ss = Samplesheet(path)
        self.samplesheet = self.ss.sample_table
        self.caption_libs = 'Mean absolute difference of percentile ranking between PGSs estimated from imputed genotyping data of eight genotyping arrays and six LPS coverages and PGS estimated from WGS in 5 different populations with PRsice p-value setting of '


    def rendering(self):
        ## Render percentile
        for row in self.samplesheet.to_dict(orient='records'):
            caption = f'Figure S. {row["index"]} ' + self.caption_libs + row['percentile']
            self.addPicture(row['file'])
            self.addCaption(caption=caption)
            # break page
            self.docx.add_page_break()
        
        ## Save to file
        self.save()

    def addCaption(self, caption: str):
        p = self.docx.add_paragraph()
        p.add_run(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def addPicture(self, picture_path: str):
        # Create temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
            temp_image_path = tmp_img.name
        
        # Convert PDF page to image
        doc = fitz.open(picture_path)
        page = doc.load_page(0)  # First page
        pix = page.get_pixmap(dpi=200)
        pix.save(temp_image_path)
        doc.close()

        # add picture
        self.docx.add_picture(temp_image_path, width=Inches(6))

        # Delete the temporary image
        os.remove(temp_image_path)

    def save(self) -> None:
        self.docx.save(self.docx_path)

            
    

if __name__ == '__main__':
    path = "./"
    ss = Samplesheet(path)

    # print(ss.listPdfFiles())
    # print(ss.sample_table)

    ## Render words
    # pp = PdfPrinter('test.docx')

    # pp.addTable("/home/ktest/project/truongphi/RND/lps_paper/out_tables/Tab_S3_mean_r2_summary_bin_(0.05–0.5].csv")

    # pp.save()

    ## Render all
    pp = PdfPrinter('supplements.docx', path)
    pp.rendering()

    